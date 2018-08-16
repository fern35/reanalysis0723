from sklearn.model_selection import train_test_split
from sklearn.model_selection import GridSearchCV, cross_val_score
from sklearn.ensemble import RandomForestRegressor,GradientBoostingRegressor
from sklearn.metrics import mean_squared_error
from docx import Document
import os
import numpy as np


class Modeler(object):

    def __init__(self, datestr, datasavedir=""):
        if datasavedir == "":
            self.basedir = os.path.abspath(os.path.dirname(__file__))
            self.datasavedir = os.path.abspath(
                os.path.join(self.basedir, '..','..','data_save{}'.format(datestr)))
        else:
            self.datasavedir = datasavedir


    def check_savepath(self,foldpath, filename):
        if not os.path.exists(foldpath):
            os.makedirs(foldpath)
        save_path = os.path.join(foldpath,filename)
        try:
            os.remove(save_path)
        except OSError:
            pass
        return save_path

    def train_RandomForest(self,X,y,title):
        X_train, X_test, y_train, y_test = train_test_split(
            X, y, test_size=0.3, random_state=0)
        print(X_train[:10])
        print("==========")
        print(y_train[:10])
        # Set the parameters by cross-validation
        tuned_parameters = [{'n_estimators': [10,20,30],
                             'max_features': ['auto','sqrt',0.2,0.4,0.6],
                             'min_samples_split': [2,3, 10, 100, 1000]}]

        scores = ['neg_mean_squared_error']

        for score in scores:
            document = Document()
            document.add_heading('{}_RandomForest_{}'.format(title,score), 0)

            document.add_paragraph("# Tuning hyper-parameters for %s" % score)

            clf = GridSearchCV(RandomForestRegressor(), tuned_parameters, cv=5,
                               scoring=score)
            clf.fit(X_train, y_train)

            document.add_paragraph("Best parameters set found on development set:")
            document.add_paragraph(clf.best_params_)

            document.add_paragraph("Grid scores on development set:")
            means = clf.cv_results_['mean_test_score']
            stds = clf.cv_results_['std_test_score']
            for mean, std, params in zip(means, stds, clf.cv_results_['params']):
                document.add_paragraph("%0.3f (+/-%0.03f) for %r"
                      % (mean, std * 2, params))

            document.add_paragraph()
            document.add_paragraph("The scores are computed on the full evaluation set")

            y_true, y_pred = np.squeeze(y_test).tolist(), clf.predict(X_test).tolist()

            document.add_paragraph(str(mean_squared_error(y_true, y_pred)))
            fold_path = os.path.join(self.datasavedir, 'doc', 'Model')
            save_path = self.check_savepath(foldpath=fold_path, filename='{}_RandomForest_{}.docx'.format(title,score))
            document.save(save_path)

    def train_GradientBoosting(self,X,y,title):
        X_train, X_test, y_train, y_test = train_test_split(
            X, y, test_size=0.3, random_state=0)

        # Set the parameters by cross-validation
        tuned_parameters = [{'learning_rate': [0.01,0.05,0.1],
                             'n_estimators':[100,200,300],
                             'max_features': ['auto','sqrt',0.2,0.4,0.6]
                             }]

        scores = ['neg_mean_squared_error']

        for score in scores:
            document = Document()
            document.add_heading('{}_GradientBoosting_{}'.format(title,score), 0)

            document.add_paragraph("# Tuning hyper-parameters for %s" % score)

            clf = GridSearchCV(GradientBoostingRegressor(), tuned_parameters, cv=5,
                               scoring=score)
            clf.fit(X_train, y_train)

            document.add_paragraph("Best parameters set found on development set:")
            document.add_paragraph(clf.best_params_)

            document.add_paragraph("Grid scores on development set:")
            means = clf.cv_results_['mean_test_score']
            stds = clf.cv_results_['std_test_score']
            for mean, std, params in zip(means, stds, clf.cv_results_['params']):
                document.add_paragraph("%0.3f (+/-%0.03f) for %r"
                      % (mean, std * 2, params))

            document.add_paragraph()
            document.add_paragraph("The scores are computed on the full evaluation set")

            y_true, y_pred = np.squeeze(y_test).tolist(), clf.predict(X_test).tolist()

            document.add_paragraph(str(mean_squared_error(y_true, y_pred)))
            fold_path = os.path.join(self.datasavedir, 'doc', 'Model')
            save_path = self.check_savepath(foldpath=fold_path, filename='{}_GradientBoosting_{}.docx'.format(title,score))
            document.save(save_path)
