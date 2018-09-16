import os
from docx import Document
import pandas as pd
import matplotlib.pyplot as plt
from docx.shared import Inches
import seaborn as sns
from utils.constants import Armoire_GROUP,Armoire_PICK
from utils.constants import PL_GROUP,PL_PICK
from utils.constants import Int_GROUP,Int_PICK
from utils.constants import VILLE_NAME
from modules.saver import Saver
from sklearn.cluster import KMeans
import numpy as np
import re
import string
import seaborn as sns


class Analyzer(object):
    def __init__(self, datestr, datasavedir=""):
        if datasavedir == "":
            self.basedir = os.path.abspath(os.path.dirname(__file__))
            self.datasavedir = os.path.abspath(
                os.path.join(self.basedir, '..','..','data_save{}'.format(datestr)))
        else:
            self.datasavedir = datasavedir
        self.saver = Saver(datestr=datestr)

    def pick_Var(self,data, Var_lst):
        new_data = data[Var_lst]
        return new_data

    def check_savepath(self,foldpath, filename):
        if not os.path.exists(foldpath):
            os.makedirs(foldpath)
        save_path = os.path.join(foldpath,filename)
        try:
            os.remove(save_path)
        except OSError:
            pass
        return save_path

    def gen_NAN_excel(self,data,foldername='',titlename='',save=True):
        """generate data frame describing the condition of NAN for variables in one file"""
        stat_data = pd.DataFrame(columns=['count_wnNAN'])
        stat_data['count_wnNAN'] = data.count()
        stat_data['percentage_wnNAN'] = data.count() / len(data)
        if save:
            self.saver.save_excel(data=stat_data,filename='{}_NAN.xlsx'.format(titlename),foldername=foldername)

        return stat_data

    def gen_cat_var(self, data, Var_lst):
        """
        check all the categories of the variables of the data
        :param data:
        :param Var_lst:
        :return:
        """
        result_data = pd.DataFrame(columns=Var_lst)
        for var in Var_lst:
            tp = data[var].value_counts(dropna=False)
            result_data[var] = pd.Series(tp.index)

        return result_data

    def get_tfvect_feature(self,vectorizer,save=True,foldername='',varname=''):
        df_feature = pd.DataFrame(columns=['feature','idf'])
        df_feature['feature'] = vectorizer.get_feature_names()
        def get_idf(feature):
            index_feature = vectorizer.vocabulary_[feature]
            return vectorizer.idf_[index_feature]
        df_feature['idf'] = df_feature['feature'].apply(get_idf)
        df_feature.sort_values(by='idf',inplace=True)
        if save:
            self.saver.save_excel(data=df_feature,filename='Idf_Analysis_{}.docx'.format(varname),foldername=foldername)
        return df_feature

    def ana_gapstat_kmeans(self,n_cluster_max,data_matrix,titlename,nrefs=5):

        gaps = np.zeros((len(range(1, n_cluster_max)),))
        resultsdf = pd.DataFrame({'clusterCount': [], 'gap': []})
        for gap_index, k in enumerate(range(1, n_cluster_max)):
            # Holder for reference dispersion results
            refDisps = np.zeros(nrefs)
            # For n references, generate random sample and perform kmeans getting resulting dispersion of each loop
            for i in range(nrefs):
                # Create new random reference set
                randomReference = np.random.random_sample(size=data_matrix.shape)
                # Fit to it
                km = KMeans(k)
                km.fit(randomReference)
                refDisp = km.inertia_
                refDisps[i] = refDisp
            # Fit cluster to original data and create dispersion
            km = KMeans(k)
            km.fit(data_matrix)
            origDisp = km.inertia_
            # Calculate gap statistic
            gap = np.log(np.mean(refDisps)) - np.log(origDisp)
            # Assign this loop's gap statistic to gaps
            gaps[gap_index] = gap
            resultsdf = resultsdf.append({'clusterCount': k, 'gap': gap}, ignore_index=True)

        optimalK = gaps.argmax() + 1

        plt.plot(resultsdf.clusterCount, resultsdf.gap, linewidth=3)
        plt.scatter(resultsdf[resultsdf.clusterCount == optimalK].clusterCount, resultsdf[resultsdf.clusterCount == optimalK].gap, s=250, c='r')
        plt.grid(True)
        plt.xlabel('Cluster Count')
        plt.ylabel('Gap Value')
        plt.title(titlename)
        fold_path = os.path.join(self.datasavedir, 'img', 'inertia')
        save_path = self.check_savepath(foldpath=fold_path,filename='{}.jpg'.format(titlename))
        plt.savefig(save_path)
        plt.clf()

        return optimalK, resultsdf

    def gen_groupCompl_cities(self,foldername='Armoire', villelst=VILLE_NAME,group_dict=Armoire_GROUP,threshold=0.0):
        document = Document()
        document.add_heading('{}_completeness_threshold={}'.format(foldername,threshold), 0)

        for ville in villelst:
            data = pd.read_excel(os.path.join(self.datasavedir, 'excel/{}/{}_{}.xlsx'.format(foldername,foldername,ville)))
            data_groupNAN = self.group_Compl(data, group_dict=group_dict,threshold=threshold)
            plt.clf()
            data_groupNAN.plot(kind='bar',figsize=(3,3),fontsize=8,legend=False)
            plt.title('completeness variables {}_threshold={}'.format(ville,threshold), fontsize=8)
            try:
                plt.tight_layout()
            except:
                print('Cannot use tight layout, var: ',ville)

            plt.ylabel('percentage vars >{}'.format(threshold))
            plt.savefig(os.path.abspath(os.path.join(self.datasavedir, 'img/groupCompl/{}_{}_groupCompl_threshold{}.jpg'.format(foldername,ville,threshold))))
            document.add_picture(os.path.abspath(os.path.join(self.datasavedir, 'img/groupCompl/{}_{}_groupCompl_threshold{}.jpg'.format(foldername,ville,threshold))))

        fold_path = os.path.join(self.datasavedir, 'doc', '{}'.format(foldername))
        save_path = self.check_savepath(foldpath=fold_path, filename='groupCompl{}.docx'.format(threshold))
        document.save(save_path)

    def group_Compl(self,data,group_dict,threshold):
        """generate completeness condition of each group for one city"""
        result = pd.DataFrame(index=group_dict.keys())
        freq = data.count() / len(data)
        for group,vars in group_dict.items():
            count_group = 0
            count_threshold = 0
            for var in vars:
                if var in freq.index.tolist() :
                    count_group += 1
                    if freq.loc[var]>threshold:
                        count_threshold += 1
            result.loc[group,0] = count_threshold/count_group
        return result

    def gen_VarIntersection(self,foldername='Armoire',villelst=VILLE_NAME,group_dict=Armoire_GROUP,threshold=0.0):
        result_dict = dict.fromkeys(group_dict.keys(),[])
        for ville in villelst:
            data = pd.read_excel(os.path.abspath(os.path.join(self.datasavedir, 'excel/{}/{}_{}.xlsx'.format(foldername,foldername,ville))))
            ville_dict = self.gen_VarThreshold(data,group_dict=group_dict,threshold=threshold)

            for group in group_dict:
                lst = result_dict[group].copy()
                lst.append(ville_dict[group])
                result_dict[group] = lst

        for group in group_dict:
            result_dict[group] = set.intersection(*map(set, result_dict[group]))

        # save to txt
        f = open(self.datasavedir+'/doc/{}/var_threshold{}.txt'.format(foldername,threshold), 'w')
        for key, value in result_dict.items():
            f.write(key + ':' + str(value))
            f.write('\n')
        f.close()
        return result_dict

    def gen_VarThreshold(self,data,group_dict,threshold=0.0):
        result_dict = dict.fromkeys(group_dict.keys(),[])
        stat_data = self.gen_NAN_excel(data,save=False)
        for group,vars in group_dict.items():
            lst_tp = []
            for var in vars:
                if stat_data.loc[var,'percentage_wnNAN']>threshold:
                    lst_tp.append(var)
            result_dict[group] = lst_tp
        return result_dict

    def comp_Var_cities(self,foldername='Armoire',villelst=VILLE_NAME,group_dict= Armoire_PICK):
        # first test the result of categorical variables
        document = Document()
        document.add_heading('compa_Var_{}'.format(foldername), 0)
        var_lst = group_dict['CAT']+group_dict['DIST']
        for var in var_lst:
            if var in group_dict['CAT']:
                result_cat = pd.DataFrame(columns=villelst)
                for ville in villelst:
                    data = pd.read_excel(os.path.join(self.datasavedir, 'excel','{}'.format(foldername),'{}_{}.xlsx'.format(foldername,ville)))
                    result_cat[ville] = data[var].value_counts(dropna=False)/len(data)

                # if high variety of values, split the graph
                if len(result_cat)>8:
                    no_split = int(len(result_cat)/8)
                    for i in range(no_split):
                        plt.clf()
                        result_cat.iloc[i*8:i*8+7,:].plot(kind='bar')
                        plt.title('{} for cities_BDD{}_{}:{}'.format(var, foldername,i*8,i*8+7), fontsize=12)
                        try:
                            plt.tight_layout()
                        except:
                            print('Cannot use tight layout, var: ', var)

                        plt.ylabel('percentage')
                        fold_path_mul = os.path.join(self.datasavedir,'img','varDistCities')
                        save_path_mul = self.check_savepath(foldpath=fold_path_mul,filename='{}_{}_{}:{}.jpg'.format(foldername, var,i*8,i*8+7))
                        plt.savefig(save_path_mul)
                        document.add_picture(save_path_mul,
                                             width=Inches(4.5))
                    plt.clf()
                    result_cat.iloc[no_split*8:,:].plot(kind='bar')
                    plt.title('{} for cities_BDD{}_{}:{}'.format(var, foldername,no_split*8,len(result_cat)-1), fontsize=12)
                    try:
                        plt.tight_layout()
                    except:
                        print('Cannot use tight layout, var: ', var)

                    plt.ylabel('percentage')

                    fold_path_last = os.path.join(self.datasavedir, 'img', 'varDistCities')
                    save_path_last = self.check_savepath(foldpath=fold_path_last,
                                                        filename='{}_{}_{}:{}.jpg'.format(foldername, var,no_split*8,len(result_cat)-1))
                    plt.savefig(save_path_last)
                    document.add_picture(save_path_last,width=Inches(4.5))

                else:
                    plt.clf()
                    result_cat.plot(kind='bar')
                    plt.title('{} for cities_BDD{}'.format(var,foldername), fontsize=12)
                    try:
                        plt.tight_layout()
                    except:
                        print('Cannot use tight layout, var: ',var)

                    plt.ylabel('percentage')
                    fold_path_img = os.path.join(self.datasavedir, 'img', 'varDistCities')
                    save_path_img = self.check_savepath(foldpath=fold_path_img,
                                                        filename='{}_{}.jpg'.format(foldername,var))

                    plt.savefig(save_path_img)
                    document.add_picture(save_path_img,
                                         width=Inches(4.5))
            else:
                result_dist = pd.DataFrame(columns=villelst)
                for ville in villelst:
                    data = pd.read_excel(os.path.join(self.datasavedir, 'excel','{}'.format(foldername),'{}_{}.xlsx'.format(foldername,ville)))
                    result_dist[ville] = data[var]

                plt.clf()
                result_dist.plot(kind='density')
                plt.title('{} for cities_BDD{}'.format(var, foldername), fontsize=12)
                try:
                    plt.tight_layout()
                except:
                    print('Cannot use tight layout, var: ', var)

                plt.ylabel('percentage')
                fold_path_img_var = os.path.join(self.datasavedir, 'img', 'varDistCities')
                save_path_img_var = self.check_savepath(foldpath=fold_path_img_var,filename='{}_{}.jpg'.format(foldername, var))
                plt.savefig(save_path_img_var)
                document.add_picture(save_path_img_var,
                                     width=Inches(4.5))

        fold_path_doc = os.path.join(self.datasavedir,'doc','{}'.format(foldername))
        save_path_doc = self.check_savepath(foldpath=fold_path_doc,filename='compa_Var_{}.docx'.format(foldername))
        document.save(save_path_doc)

    def gen_histogram_Pie(self,data,titlename,Var_lst):
        document = Document()

        document.add_heading('Histogram & Pie_{}_{} instances'.format(titlename,len(data)), 0)

        for var in Var_lst:
            plt.clf()
            data[var].value_counts(dropna=False).plot(kind='bar')
            plt.title('Histogram {}'.format(var), fontsize=12)
            try:
                plt.tight_layout()
            except:
                print('Cannot use tight layout, var: ',var)

            plt.ylabel('numbers')
            fold_path_bar = os.path.join(self.datasavedir,'img','{}'.format(titlename))
            save_path_bar = self.check_savepath(foldpath=fold_path_bar,filename='{}_{}_bar.jpg'.format(titlename,var))
            plt.savefig(save_path_bar)
            document.add_picture(save_path_bar,width=Inches(4.5))

            plt.clf()
            (data[var].value_counts(dropna=False)/len(data)).plot(kind='pie')
            plt.title('Pie plot {}'.format(var), fontsize=12)
            plt.tight_layout()
            fold_path_pie = os.path.join(self.datasavedir,'img','{}'.format(titlename))
            save_path_pie = self.check_savepath(foldpath=fold_path_pie,filename='{}_{}_pie.jpg'.format(titlename,var))
            plt.savefig(save_path_pie)
            document.add_picture(save_path_pie,width=Inches(4.5))

        fold_path_doc = os.path.join(self.datasavedir,'doc','{}'.format(titlename))
        save_path_doc = self.check_savepath(foldpath=fold_path_doc,filename='Histogram_Pie_{}.docx'.format(titlename))
        document.save(save_path_doc)

    def gen_Dist(self,data,titlename,Var_lst):
        document = Document()

        document.add_heading('Distribution_{}_{} instances'.format(titlename,len(data)), 0)

        for var in Var_lst:
            plt.clf()
            sns.distplot(data[var].dropna())
            plt.title('Distribution(dropna={}) {}'.format(data[var].isnull().sum(),var), fontsize=12)
            plt.ylabel('frequncy')
            plt.tight_layout()
            fold_path_img = os.path.join(self.datasavedir,'img','{}'.format(titlename))
            save_path_img = self.check_savepath(foldpath=fold_path_img,filename='{}_{}_dist.jpg'.format(titlename,var))
            plt.savefig(save_path_img)
            document.add_picture(save_path_img,width=Inches(4.5))

        fold_path_doc = os.path.join(self.datasavedir,'doc','{}'.format(titlename))
        save_path_doc = self.check_savepath(foldpath=fold_path_doc,filename='Dist_{}.docx'.format(titlename))
        document.save(save_path_doc)

    def split_cat(self,data,Var_lst):
        """
        split the multi choices into separate categories
        :param data:
        :return:
        """
        pattern = "[A-Z]"
        new_data = data.fillna('NA')
        result_data = pd.DataFrame()

        for var in Var_lst:
            cat_lst = []
            # get the total categories and corresponding occurrence
            for i in range(len(new_data)):
                old_string = new_data[var].loc[i]
                # split to list
                cat_lst_tp = old_string.split('\n')
                cat_lst_tp = [self.check_content_NL(i,pattern) for i in cat_lst_tp]

                cat_lst.extend(cat_lst_tp)

            cat_lst = list(set(cat_lst))
            print(var,', length: ',len(cat_lst))
            cat_pd = pd.DataFrame(cat_lst,columns=[var])

            result_data= pd.concat([result_data,cat_pd],axis=1)

        return result_data

    def check_content_NL(self,inner_data,pattern):
        # data : a string
        # replace the punctuation with space
        inner_new = inner_data.translate(str.maketrans(string.punctuation, ' ' * len(string.punctuation)))
        word_lst = inner_new.split()
        if any(i.isupper() for i in word_lst):
            return ' '.join([ele.lower() for ele in word_lst])
        else:
            # detect the word with uppercase beginning
            new_string = re.sub(pattern, lambda x: " " + x.group(0), inner_new).lower()
            # remove the space at the beginning and the bottom
            return ' '.join(new_string.split())

    def get_distr_inter(self,data,Var_lst,filename, foldername=''):
        """
        get the distribution of the number of interventions
        :param data: dataframe
        :param Var_lst:
        :return:
        """
        distr = data.groupby(Var_lst).size().reset_index(name='counts')
        print(distr.head())
        sns_plot = sns.distplot(distr['counts'])
        sns_plot.set( ylabel='density',title=filename)
        fold_path_img = os.path.join(self.datasavedir, 'img', '{}'.format(foldername))
        save_path_img = self.check_savepath(foldpath=fold_path_img, filename='{}_dist.jpg'.format(filename))

        plt.savefig(save_path_img)

    def split_cat_get_dist(self, data, var):

        pattern = "[A-Z]"
        new_data = data.fillna(value={var:'NA'})
        new_data.reset_index(drop=True, inplace=True)
        # result_data = pd.DataFrame()
        result_dict = {}

        # cat_lst = []
        # get the total categories and corresponding occurrence
        for i in range(len(new_data)):
            old_string = new_data[var].loc[i]
            # split to list
            cat_lst_tp = old_string.split('\n')
            cat_lst_tp = [self.check_content_NL(i, pattern) for i in cat_lst_tp]
            for cat in cat_lst_tp:
                result_dict[cat] = result_dict.get(cat, 0) + 1

        result_dict.update((x, y / len(new_data)) for x, y in result_dict.items())
        result_data = pd.Series(result_dict)
        return result_data

    def comp_Var_cluster(self,data,group_dict,n_cluster=3,foldername='Armoire'):
        """
        compare the distribution od the values of the features
        :param data: the generated dataframe for clusters
        :param group_dict: the dictionary for denoting the categories of features, example: {'CAT':[...], 'DIST':[...]}
        :param foldername:
        :return:
        """
        document = Document()
        document.add_heading('compa_Cluster_{}'.format(foldername), 0)

        var_lst = group_dict['CAT']+group_dict['DIST']+group_dict['MULTI']
        for var in var_lst:
            print('var: ',var)
            clusterindex_strlst = ['cluster_'+str(i) for i in range(n_cluster)]
            clusterindex_lst = [i for i in range(n_cluster)]
            if var in group_dict['CAT']:
                result_cat = pd.DataFrame(columns=clusterindex_strlst)
                for cluster_index in clusterindex_lst:
                    result_cat['cluster_'+str(cluster_index)] = data[data['cluster_index']==cluster_index][var].value_counts(dropna=False)/len(data)
                # self.saver.save_excel(data=result_cat,filename='result_cat_{}_ncluster{}_{}'.format(var,n_cluster,foldername),foldername=foldername)
                # if high variety of values, split the graph
                if len(result_cat)>8:
                    no_split = int(len(result_cat)/8)
                    for i in range(no_split):
                        plt.clf()
                        result_cat.iloc[i*8:i*8+7,:].plot(kind='bar')
                        plt.title('{} for {} ncluster={}_{}:{}'.format(var,n_cluster,foldername,i*8,i*8+7), fontsize=12)
                        try:
                            plt.tight_layout()
                        except:
                            print('Cannot use tight layout, var: ', var)

                        plt.ylabel('percentage')
                        fold_path_mul = os.path.join(self.datasavedir,'img','cluster_var')
                        save_path_mul = self.check_savepath(foldpath=fold_path_mul,filename='{}_ncluster{}_{}_{}:{}.jpg'.format(foldername,n_cluster,var,i*8,i*8+7))
                        plt.savefig(save_path_mul)
                        document.add_picture(save_path_mul,
                                             width=Inches(4.5))
                    plt.clf()
                    result_cat.iloc[no_split*8:,:].plot(kind='bar')
                    plt.title('{} for {} ncluster={}_{}:{}'.format(var, n_cluster,foldername,no_split*8,len(result_cat)-1), fontsize=12)
                    try:
                        plt.tight_layout()
                    except:
                        print('Cannot use tight layout, var: ', var)

                    plt.ylabel('percentage')

                    fold_path_last = os.path.join(self.datasavedir, 'img', 'cluster_var')
                    save_path_last = self.check_savepath(foldpath=fold_path_last,
                                                        filename='{}_cluster{}_{}_{}:{}.jpg'.format(foldername, n_cluster,var,no_split*8,len(result_cat)-1))
                    plt.savefig(save_path_last)
                    document.add_picture(save_path_last,width=Inches(4.5))

                else:
                    plt.clf()
                    result_cat.plot(kind='bar')
                    plt.title('{} for {} ncluster={}'.format(var,foldername,n_cluster), fontsize=12)
                    try:
                        plt.tight_layout()
                    except:
                        print('Cannot use tight layout, var: ',var)

                    plt.ylabel('percentage')
                    fold_path_img = os.path.join(self.datasavedir, 'img', 'cluster_var')
                    save_path_img = self.check_savepath(foldpath=fold_path_img,
                                                        filename='{}_cluster{}_{}.jpg'.format(foldername,n_cluster,var))

                    plt.savefig(save_path_img)
                    document.add_picture(save_path_img,
                                         width=Inches(4.5))
            elif var in group_dict['DIST']:
                result_dist = pd.DataFrame(columns=clusterindex_strlst)
                for cluster_index in clusterindex_lst:
                    data_tp = data[data['cluster_index']==cluster_index]
                    data_tp.reset_index(drop=True, inplace=True)
                    result_dist['cluster_'+str(cluster_index)] = data_tp[var]
                plt.clf()

                filtered_dist = result_dist.dropna(axis='columns', how='all')

                if len(result_dist.columns) > len(filtered_dist.columns):
                    print('Some clusters don\'t have the values for {} (n_cluster={})!'.format(var,n_cluster))
                    self.saver.save_excel(data=result_dist, filename='Strange_{}_cluster{}'.format(var, n_cluster),
                                          foldername=foldername)
                    result_dist = filtered_dist
                else:
                    pass

                result_dist.plot(kind='density')
                plt.title('{} for {} ncluster={}'.format(var, foldername, n_cluster), fontsize=12)
                try:
                    plt.tight_layout()
                except:
                    print('Cannot use tight layout, var: ', var)

                plt.ylabel('density')
                fold_path_img_var = os.path.join(self.datasavedir, 'img', 'cluster_var')
                save_path_img_var = self.check_savepath(foldpath=fold_path_img_var,
                                                        filename='{}_cluster{}_{}.jpg'.format(foldername, n_cluster,
                                                                                              var))
                plt.savefig(save_path_img_var)
                document.add_picture(save_path_img_var,
                                     width=Inches(4.5))
            else: # 'MULTI'
                result_multi = pd.DataFrame(columns=clusterindex_strlst)
                for cluster_index in clusterindex_lst:
                    result_multi['cluster_'+str(cluster_index)] = self.split_cat_get_dist(data=data[data['cluster_index']==cluster_index],var=var)

                plt.clf()
                result_multi.plot(kind='bar')
                plt.title('{} for {} ncluster={}'.format(var, foldername, n_cluster), fontsize=12)
                try:
                    plt.tight_layout()
                except:
                    print('Cannot use tight layout, var: ', var)

                plt.ylabel('percentage')
                fold_path_img = os.path.join(self.datasavedir, 'img', 'cluster_var')
                save_path_img = self.check_savepath(foldpath=fold_path_img,
                                                    filename='{}_cluster{}_{}.jpg'.format(foldername, n_cluster, var))

                plt.savefig(save_path_img)
                document.add_picture(save_path_img,
                                     width=Inches(4.5))


        fold_path_doc = os.path.join(self.datasavedir,'doc','{}'.format(foldername))
        save_path_doc = self.check_savepath(foldpath=fold_path_doc,filename='compa_cluster{}_{}.docx'.format(n_cluster,foldername))
        document.save(save_path_doc)

    def plot_feature_importance(self,importances,featurenames,title,top_n=30):
        indices = np.argsort(importances)[::-1]
        plt.clf()
        plt.figure()
        plt.title(title)
        feature_plot = featurenames[indices][:top_n]
        print(feature_plot)

        plt.barh(range(top_n), importances[indices][:top_n],
                color="r",
                # yerr=std[indices],
                align="center")
        plt.yticks(range(top_n), feature_plot[:top_n])
        plt.ylim([-1, top_n])
        plt.tight_layout()
        fold_path = os.path.join(self.datasavedir,'img','model')
        save_path = self.check_savepath(foldpath=fold_path,filename='{}.jpg'.format(title))
        plt.savefig(save_path)
        plt.show()


